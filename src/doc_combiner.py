import os
import sys
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QPushButton, QVBoxLayout, 
                            QWidget, QLabel, QFileDialog, QListWidget, QMessageBox,
                            QHBoxLayout, QLineEdit, QInputDialog, QListWidgetItem)
from PyQt6.QtCore import Qt
from docx import Document
import docx2txt
import sys

class DocCombinerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Combiner")
        self.setMinimumSize(800, 600)
        
        # Initialize variables
        self.files = []
        self.output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'results')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Create main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout()
        
        # Output directory display
        output_layout = QHBoxLayout()
        output_layout.addWidget(QLabel("Output Directory:"))
        self.output_dir_label = QLineEdit(self.output_dir)
        self.output_dir_label.setReadOnly(True)
        output_layout.addWidget(self.output_dir_label)
        btn_browse_output = QPushButton("Change")
        btn_browse_output.clicked.connect(self.change_output_dir)
        output_layout.addWidget(btn_browse_output)
        layout.addLayout(output_layout)
        
        # Add files section
        btn_add_files = QPushButton("Add Word Documents")
        btn_add_files.clicked.connect(self.add_files)
        layout.addWidget(btn_add_files)
        
        # Files list with roles
        self.files_list = QListWidget()
        self.files_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        layout.addWidget(QLabel("Files to process (double-click to edit role):"))
        layout.addWidget(self.files_list)
        
        # Button to remove selected files
        btn_remove = QPushButton("Remove Selected")
        btn_remove.clicked.connect(self.remove_selected_files)
        layout.addWidget(btn_remove)
        
        # Process button
        btn_process = QPushButton("Process and Combine Documents")
        btn_process.clicked.connect(self.process_documents)
        layout.addWidget(btn_process)
        
        main_widget.setLayout(layout)
        
        # Connect double-click event to edit role
        self.files_list.itemDoubleClicked.connect(self.edit_role)
    
    def add_files(self):
        file_dialog = QFileDialog()
        files, _ = file_dialog.getOpenFileNames(
            self,
            "Select Word Documents",
            "",
            "Word Documents (*.docx *.doc);;All Files (*)"
        )
        
        for file_path in files:
            if file_path not in [f['path'] for f in self.files]:
                role = os.path.splitext(os.path.basename(file_path))[0]
                self.files.append({"path": file_path, "role": role})
                self.files_list.addItem(f"{role} - {os.path.basename(file_path)}")
    
    def remove_selected_files(self):
        selected_items = self.files_list.selectedItems()
        if not selected_items:
            return
            
        for item in selected_items:
            index = self.files_list.row(item)
            self.files_list.takeItem(index)
            self.files.pop(index)
    
    def edit_role(self, item):
        index = self.files_list.row(item)
        current_role = self.files[index]['role']
        
        new_role, ok = QInputDialog.getText(
            self, 
            "Edit Role", 
            "Enter role/title for this document:", 
            text=current_role
        )
        
        if ok and new_role.strip():
            self.files[index]['role'] = new_role.strip()
            self.update_files_list()
    
    def update_files_list(self):
        self.files_list.clear()
        for file_info in self.files:
            self.files_list.addItem(f"{file_info['role']} - {os.path.basename(file_info['path'])}")
    
    def change_output_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Select Output Directory")
        if dir_path:
            self.output_dir = dir_path
            self.output_dir_label.setText(dir_path)
    
    def process_documents(self):
        if not self.files:
            QMessageBox.warning(self, "No Files", "Please add files to process.")
            return
        
        output_doc = Document()
        
        for file_info in self.files:
            file_path = file_info['path']
            role = file_info['role']
            
            try:
                # Add role as heading with some formatting
                heading = output_doc.add_heading(level=1)
                run = heading.add_run(role)
                run.bold = True
                
                # Extract text from the document
                text = docx2txt.process(file_path)
                
                # Add the extracted text with basic formatting
                if text.strip():
                    output_doc.add_paragraph(text)
                
                # Add a page break between documents
                output_doc.add_paragraph()
                output_doc.add_page_break()
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error processing {file_path}: {str(e)}")
                return
        
        # Save the combined document
        output_path = os.path.join(self.output_dir, "combined_documents.docx")
        counter = 1
        while os.path.exists(output_path):
            output_path = os.path.join(self.output_dir, f"combined_documents_{counter}.docx")
            counter += 1
        
        output_doc.save(output_path)
        
        QMessageBox.information(
            self, 
            "Success", 
            f"Documents combined successfully!\nSaved to: {output_path}"
        )
        
        # Open the output directory
        os.startfile(os.path.dirname(output_path))

def main():
    app = QApplication(sys.argv)
    window = DocCombinerApp()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
